"""Resume document generation — PDF via HTML + WeasyPrint, DOCX via python-docx."""
import html as _html


def _esc(text: str) -> str:
    if not text:
        return ""
    return _html.escape(str(text), quote=True)


# ═══════════════════════════════════════════════════════════════
# Shared CSS — clean, simple, with proper page margins for PDF
# ═══════════════════════════════════════════════════════════════

_RESUME_CSS = """
  @page { size: A4; margin: 1.8cm 2cm 1.8cm 2cm; }

  body {
    font-family: "WenQuanYi Micro Hei", "Microsoft YaHei", "PingFang SC", sans-serif;
    color: #222; font-size: 10pt; line-height: 1.55;
    word-break: break-word; margin: 0; padding: 0;
  }

  /* ── Header ── */
  .header { margin-bottom: 16px; }
  .header-table { width: 100%; border-collapse: collapse; }
  .header-table td { vertical-align: middle; }
  .header-table td.info-col { text-align: left; }
  .header-table td.photo-col { width: 90px; text-align: right; }
  .header-name {
    font-size: 20pt; font-weight: 700; color: #1a1a1a;
    margin: 0 0 4px 0; letter-spacing: 2px;
  }
  .header-contact {
    font-size: 9pt; color: #555; margin: 0;
  }
  .header-contact span { margin-right: 12px; }
  .photo-box {
    width: 80px; height: 100px; border: 1px solid #ccc;
    text-align: center; vertical-align: middle; display: table-cell;
    color: #999; font-size: 7.5pt;
  }

  /* ── Section ── */
  .section { margin-bottom: 12px; page-break-inside: avoid; }
  .section h2 {
    font-size: 11pt; font-weight: 700; color: #333;
    padding-bottom: 3px; border-bottom: 1px solid #999;
    margin: 0 0 6px 0;
  }

  /* ── Target ── */
  .target-line {
    font-size: 10pt; color: #333; margin-bottom: 12px; padding: 0;
  }
  .target-line strong { font-weight: 600; }

  /* ── Summary ── */
  .summary-text {
    font-size: 9.5pt; color: #333; line-height: 1.6;
    margin: 0 0 12px 0; padding: 0;
  }

  /* ── Items ── */
  .item { margin-bottom: 10px; page-break-inside: avoid; }
  .item-header {
    font-size: 10.5pt; font-weight: 700; color: #1a1a1a; margin: 0;
  }
  .item-sub {
    font-size: 9pt; color: #555; margin: 0 0 2px 0;
  }
  .item p {
    font-size: 9.5pt; color: #333; margin: 2px 0; line-height: 1.5;
  }
  .item ul { margin: 2px 0 0 0; padding-left: 16px; }
  .item li {
    font-size: 9pt; color: #444; margin-bottom: 1px; line-height: 1.4;
  }

  /* ── Skills ── */
  .skills-table { width: 100%; border-collapse: collapse; }
  .skills-table td { padding: 1px 8px 1px 0; font-size: 9pt; vertical-align: top; }
  .skills-table td.cat { width: 80px; font-weight: 600; color: #333; }
  .skills-table td.items { color: #444; }

  /* ── Footer ── */
  .section-footer {
    text-align: center; margin-top: 16px;
    color: #bbb; font-size: 7.5pt;
  }
"""


# ═══════════════════════════════════════════════════════════════
# HTML / PDF
# ═══════════════════════════════════════════════════════════════

def build_html_resume(resume_data: dict, template: str = "professional") -> str:
    """Build a clean, simple resume HTML string."""
    basic = resume_data.get("basic_info") or {}
    skills = resume_data.get("skills") or []
    education = resume_data.get("education") or []
    experiences = resume_data.get("work_experiences") or resume_data.get("work_experience") or []
    projects = resume_data.get("projects") or []
    summary = resume_data.get("summary") or ""
    target = resume_data.get("target_position") or (resume_data.get("job_target") or {}).get("target_position") or ""
    esc = _esc

    name = esc(basic.get("name", "")) or "姓名"
    email = esc(basic.get("email", ""))
    phone = esc(basic.get("phone", ""))
    city = esc(basic.get("city") or basic.get("address") or "")
    contact_parts = "".join(f"<span>{c}</span>" for c in [email, phone, city] if c)

    lines = []
    w = lines.append

    # ── Header (left info + right photo) ──
    w(f'<div class="header">')
    w(f'<table class="header-table"><tr>')
    w(f'<td class="info-col">')
    w(f'<p class="header-name">{name}</p>')
    if contact_parts:
        w(f'<p class="header-contact">{contact_parts}</p>')
    w(f'</td>')
    w(f'<td class="photo-col"><div class="photo-box">照片粘贴处</div></td>')
    w(f'</tr></table>')
    w(f'</div>')

    # ── Target ──
    if target:
        w(f'<div class="target-line"><strong>求职目标：</strong>{esc(target)}</div>')

    # ── Summary ──
    if summary:
        w(f'<div class="section">')
        w(f'<h2>自我评价</h2>')
        w(f'<p class="summary-text">{esc(summary)}</p>')
        w(f'</div>')

    # ── Work Experience ──
    if experiences:
        w(f'<div class="section">')
        w(f'<h2>工作经历</h2>')
        for exp in experiences:
            co, ti = esc(exp.get("company", "")), esc(exp.get("title", ""))
            loc = esc(exp.get("location", ""))
            st = esc(str(exp.get("start_date", "")))
            en = esc(str(exp.get("end_date", "")) or "至今")
            head = f"{co} — {ti}" if ti else co
            sub = f"{loc} · {st} ~ {en}" if loc else f"{st} ~ {en}"
            w(f'<div class="item">')
            w(f'<p class="item-header">{head}</p>')
            w(f'<p class="item-sub">{sub}</p>')
            if exp.get("description"):
                w(f'<p>{esc(exp["description"])}</p>')
            highlights = exp.get("highlights") or []
            if highlights:
                w(f'<ul>')
                for h in highlights:
                    w(f'<li>{esc(str(h))}</li>')
                w(f'</ul>')
            w(f'</div>')
        w(f'</div>')

    # ── Projects ──
    if projects:
        w(f'<div class="section">')
        w(f'<h2>项目经历</h2>')
        for proj in projects:
            nm = esc(proj.get("name", ""))
            role = esc(proj.get("role", ""))
            st = esc(str(proj.get("start_date", "")))
            en = esc(str(proj.get("end_date", "")))
            sp = " · ".join(x for x in [role, f"{st} ~ {en}".strip(" ~") if st or en else ""] if x)
            w(f'<div class="item">')
            w(f'<p class="item-header">{nm}</p>')
            if sp:
                w(f'<p class="item-sub">{sp}</p>')
            tech = proj.get("tech_stack") or []
            if tech:
                w(f'<p>{", ".join(esc(str(t)) for t in tech)}</p>')
            if proj.get("description"):
                w(f'<p>{esc(proj["description"])}</p>')
            highlights = proj.get("highlights") or []
            if highlights:
                w(f'<ul>')
                for h in highlights:
                    w(f'<li>{esc(str(h))}</li>')
                w(f'</ul>')
            w(f'</div>')
        w(f'</div>')

    # ── Education ──
    if education:
        w(f'<div class="section">')
        w(f'<h2>教育背景</h2>')
        for edu in education:
            sc, mj = esc(edu.get("school", "")), esc(edu.get("major", ""))
            deg = esc(edu.get("degree", ""))
            st, en = esc(str(edu.get("start_date", ""))), esc(str(edu.get("end_date", "")))
            head = f"{sc} — {mj}" if mj else sc
            sub_parts = [x for x in [deg, f"{st} ~ {en}".strip(" ~") if st or en else ""] if x]
            w(f'<div class="item">')
            w(f'<p class="item-header">{head}</p>')
            if sub_parts:
                w(f'<p class="item-sub">{" · ".join(sub_parts)}</p>')
            if edu.get("description"):
                w(f'<p>{esc(edu["description"])}</p>')
            w(f'</div>')
        w(f'</div>')

    # ── Skills ──
    if skills:
        groups: dict[str, list[str]] = {}
        for s in skills:
            cat = s.get("category", "其他")
            label = esc(s.get("name", ""))
            lvl = esc(s.get("level", ""))
            if lvl:
                label += f" ({lvl})"
            groups.setdefault(cat, []).append(label)
        w(f'<div class="section">')
        w(f'<h2>专业技能</h2>')
        w(f'<table class="skills-table">')
        for cat, names in groups.items():
            w(f'<tr><td class="cat">{esc(cat)}</td><td class="items">{", ".join(names)}</td></tr>')
        w(f'</table>')
        w(f'</div>')

    # ── Footer ──
    w(f'<div class="section-footer">AI Career Copilot · 专业简历</div>')

    body = "\n".join(lines)

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<style>{_RESUME_CSS}</style>
</head>
<body>
<div class="resume-container">
{body}
</div>
</body>
</html>"""


# ═══════════════════════════════════════════════════════════════
# DOCX — clean, simple, standard page margins, no dark styling
# ═══════════════════════════════════════════════════════════════

def build_docx_resume(resume_data: dict) -> bytes:
    """Build a clean DOCX resume — standard margins, black text, no dark backgrounds."""
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    basic = resume_data.get("basic_info") or {}
    skills = resume_data.get("skills") or []
    education = resume_data.get("education") or []
    experiences = resume_data.get("work_experiences") or resume_data.get("work_experience") or []
    projects = resume_data.get("projects") or []
    summary = resume_data.get("summary") or ""
    target = resume_data.get("target_position") or (resume_data.get("job_target") or {}).get("target_position") or ""

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.15

    # ── Helpers ──
    def _section_heading(text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        # bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="999999"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def _item_header(text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10.5)

    def _item_sub(text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    def _body_text(text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(str(text))
        run.font.size = Pt(9.5)

    def _bullet(text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.6)
        run = p.add_run("• " + str(text))
        run.font.size = Pt(9)

    # ── HEADER (left info + right photo) ──
    header_tbl = doc.add_table(rows=1, cols=2)
    header_tbl.autofit = True
    header_tbl.columns[0].width = Cm(14)
    header_tbl.columns[1].width = Cm(3.2)

    # Remove table borders
    tblPr = header_tbl._tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        header_tbl._tbl.insert(0, tblPr)
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

    left_cell = header_tbl.cell(0, 0)
    name_p = left_cell.paragraphs[0]
    name_p.paragraph_format.space_after = Pt(4)
    name_run = name_p.add_run(basic.get("name", "姓名"))
    name_run.bold = True
    name_run.font.size = Pt(20)

    contact_bits = [x for x in [basic.get("email", ""), basic.get("phone", ""),
                                basic.get("city") or basic.get("address", "")] if x]
    if contact_bits:
        cp = left_cell.add_paragraph()
        cp.paragraph_format.space_after = Pt(0)
        cr = cp.add_run("  |  ".join(contact_bits))
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    right_cell = header_tbl.cell(0, 1)
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add photo border via cell shading
    tcPr = right_cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    rr1 = rp.add_run("照片")
    rr1.font.size = Pt(7.5)
    rr1.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    rp2 = right_cell.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr2 = rp2.add_run("粘贴处")
    rr2.font.size = Pt(7.5)
    rr2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Spacing after header
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

    # ── TARGET ──
    if target:
        tp = doc.add_paragraph()
        tp.paragraph_format.space_after = Pt(10)
        tr1 = tp.add_run("求职目标：")
        tr1.bold = True
        tr1.font.size = Pt(10)
        tr2 = tp.add_run(target)
        tr2.font.size = Pt(10)

    # ── SUMMARY ──
    if summary:
        _section_heading("自我评价")
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(8)
        sr = sp.add_run(summary)
        sr.font.size = Pt(9.5)

    # ── WORK EXPERIENCE ──
    if experiences:
        _section_heading("工作经历")
        for exp in experiences:
            co = exp.get("company", "")
            ti = exp.get("title", "")
            loc = exp.get("location", "")
            st = str(exp.get("start_date", ""))
            en = str(exp.get("end_date", "")) or "至今"
            head = f"{co} — {ti}" if ti else co
            sub = f"{loc} · {st} ~ {en}" if loc else f"{st} ~ {en}"
            _item_header(head)
            _item_sub(sub)
            if exp.get("description"):
                _body_text(exp["description"])
            for h in (exp.get("highlights") or []):
                _bullet(str(h))

    # ── PROJECTS ──
    if projects:
        _section_heading("项目经历")
        for proj in projects:
            _item_header(proj.get("name", ""))
            role = proj.get("role", "")
            st = str(proj.get("start_date", ""))
            en = str(proj.get("end_date", ""))
            sub_parts = [x for x in [role, f"{st} ~ {en}".strip(" ~") if st or en else ""] if x]
            if sub_parts:
                _item_sub(" · ".join(sub_parts))
            tech = proj.get("tech_stack") or []
            if tech:
                _body_text(", ".join(str(t) for t in tech))
            if proj.get("description"):
                _body_text(proj["description"])
            for h in (proj.get("highlights") or []):
                _bullet(str(h))

    # ── EDUCATION ──
    if education:
        _section_heading("教育背景")
        for edu in education:
            sc = edu.get("school", "")
            mj = edu.get("major", "")
            head = f"{sc} — {mj}" if mj else sc
            _item_header(head)
            deg = edu.get("degree", "")
            st = str(edu.get("start_date", ""))
            en = str(edu.get("end_date", ""))
            sub_parts = [x for x in [deg, f"{st} ~ {en}".strip(" ~") if st or en else ""] if x]
            if sub_parts:
                _item_sub(" · ".join(sub_parts))
            if edu.get("description"):
                _body_text(edu["description"])

    # ── SKILLS ──
    if skills:
        _section_heading("专业技能")
        groups: dict[str, list[str]] = {}
        for s in skills:
            cat = s.get("category", "其他")
            label = s.get("name", "")
            lvl = s.get("level", "")
            if lvl:
                label += f" ({lvl})"
            groups.setdefault(cat, []).append(label)
        for cat, names in groups.items():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            rc = p.add_run(f"{cat}: ")
            rc.bold = True
            rc.font.size = Pt(9)
            p.add_run(", ".join(names)).font.size = Pt(9)

    # ── FOOTER ──
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(20)
    fr = fp.add_run("AI Career Copilot · 专业简历")
    fr.font.size = Pt(7.5)
    fr.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
